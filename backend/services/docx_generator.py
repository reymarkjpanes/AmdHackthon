"""
DOCX Report Generator for Clausify AI.
Generates professionally-designed Word documents from AnalysisResult data.
Uses a clean slate/dark color palette with AMD red accents.
"""

import io
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from models.response import AnalysisResult

logger = logging.getLogger(__name__)

# Color Palette
AMD_RED = RGBColor(0xED, 0x1C, 0x24)
SLATE_900 = RGBColor(0x1E, 0x29, 0x3B)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
GREEN_600 = RGBColor(0x16, 0xA3, 0x4A)
AMBER_600 = RGBColor(0xD9, 0x77, 0x06)
BLUE_700 = RGBColor(0x1D, 0x4E, 0xD8)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shading_elm)


class DOCXGenerator:
    """Generates modern DOCX reports from AnalysisResult data."""

    def generate_report(self, analysis: AnalysisResult, session_id: str) -> bytes:
        """Generate a complete DOCX report."""
        doc = Document()
        today = datetime.utcnow().strftime("%B %d, %Y")

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = SLATE_700

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        self._title_page(doc, today, session_id)
        self._analytics_dashboard(doc, analysis)
        self._executive_summary(doc, analysis)
        self._risk_analysis(doc, analysis)
        if analysis.comparisonMatrix:
            self._comparison_matrix(doc, analysis)
        if analysis.conflicts:
            self._conflicts(doc, analysis)
        self._recommendation(doc, analysis)
        self._footer(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        logger.info(f"DOCX generated for session {session_id} ({len(docx_bytes)} bytes)")
        return docx_bytes

    def _title_page(self, doc: Document, today: str, session_id: str):
        """Build title page."""
        for _ in range(3):
            doc.add_paragraph()

        # Title
        title = doc.add_heading("CLAUSIFY AI", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(34)
            run.font.color.rgb = SLATE_900

        # Subtitle
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run("Document Intelligence Report")
        run.font.size = Pt(14)
        run.font.color.rgb = SLATE_500

        doc.add_paragraph()

        # AMD line
        amd = doc.add_paragraph()
        amd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = amd.add_run("Powered by AMD Instinct MI300X")
        run.font.size = Pt(11)
        run.font.color.rgb = AMD_RED
        run.bold = True

        doc.add_paragraph()
        doc.add_paragraph()

        # Metadata
        meta_items = [
            ("Report Date", today),
            ("Session", session_id[:8] + "..." if len(session_id) > 8 else session_id),
            ("Platform", "Clausify AI v1.0"),
        ]
        for label, value in meta_items:
            para = doc.add_paragraph()
            r1 = para.add_run(f"{label}:  ")
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = SLATE_500
            r2 = para.add_run(value)
            r2.font.size = Pt(10)
            r2.font.color.rgb = SLATE_900

        doc.add_page_break()

    def _section_heading(self, doc: Document, text: str):
        """Add a styled section heading."""
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.color.rgb = SLATE_900
            run.font.size = Pt(16)

    def _analytics_dashboard(self, doc: Document, analysis: AnalysisResult):
        """Build analytics dashboard with key metrics."""
        self._section_heading(doc, "Analytics Dashboard")

        total_risks = len(analysis.risks)
        high = len([r for r in analysis.risks if r.level == "HIGH"])
        med = len([r for r in analysis.risks if r.level == "MEDIUM"])
        low = len([r for r in analysis.risks if r.level == "LOW"])
        conflicts = len(analysis.conflicts)
        confidence = int(analysis.recommendation.confidence * 100)

        # Metrics table
        table = doc.add_table(rows=2, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["TOTAL RISKS", "HIGH SEVERITY", "CONFLICTS", "AI CONFIDENCE"]
        values = [str(total_risks), str(high), str(conflicts), f"{confidence}%"]
        value_colors = [
            AMD_RED if high > 0 else GREEN_600,
            AMD_RED if high > 0 else GREEN_600,
            AMD_RED if conflicts > 0 else GREEN_600,
            GREEN_600 if confidence >= 70 else AMBER_600 if confidence >= 40 else AMD_RED,
        ]

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            _set_cell_bg(cell, "F1F5F9")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header)
            run.font.size = Pt(8)
            run.font.color.rgb = SLATE_500
            run.bold = True

        for i, (value, color) in enumerate(zip(values, value_colors)):
            cell = table.rows[1].cells[i]
            cell.text = ""
            _set_cell_bg(cell, "F1F5F9")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = color

        doc.add_paragraph()

        # Risk breakdown
        if total_risks > 0:
            breakdown = doc.add_paragraph()
            breakdown.add_run("Risk Breakdown:  ").bold = True
            if high > 0:
                r = breakdown.add_run(f"{high} HIGH")
                r.bold = True
                r.font.color.rgb = AMD_RED
                breakdown.add_run("  |  ")
            if med > 0:
                r = breakdown.add_run(f"{med} MEDIUM")
                r.bold = True
                r.font.color.rgb = AMBER_600
                breakdown.add_run("  |  ")
            if low > 0:
                r = breakdown.add_run(f"{low} LOW")
                r.bold = True
                r.font.color.rgb = GREEN_600

        categories = sorted(set(r.category for r in analysis.risks))
        if categories:
            cat_p = doc.add_paragraph()
            cat_p.add_run("Categories: ").bold = True
            cat_p.add_run(", ".join(categories))

        doc.add_paragraph()

    def _executive_summary(self, doc: Document, analysis: AnalysisResult):
        """Build executive summary."""
        self._section_heading(doc, "Executive Summary")
        para = doc.add_paragraph(analysis.executiveSummary)
        para.paragraph_format.line_spacing = 1.5
        doc.add_paragraph()

    def _risk_analysis(self, doc: Document, analysis: AnalysisResult):
        """Build risk analysis table."""
        self._section_heading(doc, "Risk Analysis")

        if not analysis.risks:
            p = doc.add_paragraph()
            r = p.add_run("No risks identified - all documents passed analysis.")
            r.font.color.rgb = GREEN_600
            doc.add_paragraph()
            return

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header
        headers = ["SEVERITY", "CATEGORY", "DESCRIPTION", "SOURCE"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            _set_cell_bg(cell, "1E293B")
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE_RGB

        level_colors = {"HIGH": AMD_RED, "MEDIUM": AMBER_600, "LOW": GREEN_600}
        level_bg = {"HIGH": "FEF2F2", "MEDIUM": "FFFBEB", "LOW": "F0FDF4"}

        for risk in analysis.risks:
            row = table.add_row()

            # Severity
            cell = row.cells[0]
            cell.text = ""
            _set_cell_bg(cell, level_bg.get(risk.level, "FFFFFF"))
            run = cell.paragraphs[0].add_run(risk.level)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = level_colors.get(risk.level, SLATE_700)

            # Category
            row.cells[1].text = ""
            _set_cell_bg(row.cells[1], level_bg.get(risk.level, "FFFFFF"))
            r = row.cells[1].paragraphs[0].add_run(risk.category)
            r.font.size = Pt(9)

            # Description
            row.cells[2].text = ""
            _set_cell_bg(row.cells[2], level_bg.get(risk.level, "FFFFFF"))
            r = row.cells[2].paragraphs[0].add_run(risk.description)
            r.font.size = Pt(9)

            # Source
            row.cells[3].text = ""
            _set_cell_bg(row.cells[3], level_bg.get(risk.level, "FFFFFF"))
            r = row.cells[3].paragraphs[0].add_run(risk.sourceDocument)
            r.font.size = Pt(8)
            r.font.color.rgb = SLATE_500

        doc.add_paragraph()

    def _comparison_matrix(self, doc: Document, analysis: AnalysisResult):
        """Build comparison matrix table."""
        self._section_heading(doc, "Document Comparison")

        all_cols: list[str] = []
        for row in analysis.comparisonMatrix:
            for col in row.values.keys():
                if col not in all_cols:
                    all_cols.append(col)

        num_cols = len(all_cols) + 2
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = "Table Grid"

        # Header
        header_texts = ["CRITERIA"] + [c.upper() for c in all_cols] + ["BEST"]
        for i, text in enumerate(header_texts):
            cell = table.rows[0].cells[i]
            cell.text = ""
            _set_cell_bg(cell, "1E293B")
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE_RGB

        for matrix_row in analysis.comparisonMatrix:
            row = table.add_row()
            # Criteria
            row.cells[0].text = ""
            r = row.cells[0].paragraphs[0].add_run(matrix_row.field)
            r.bold = True
            r.font.size = Pt(9)

            # Values
            for j, col in enumerate(all_cols, start=1):
                val = matrix_row.values.get(col, "-")
                is_winner = matrix_row.winner == col
                cell = row.cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                if is_winner:
                    run.bold = True
                    run.font.color.rgb = GREEN_600

            # Winner
            winner_cell = row.cells[num_cols - 1]
            winner_cell.text = ""
            run = winner_cell.paragraphs[0].add_run(matrix_row.winner or "-")
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = GREEN_600

        doc.add_paragraph()

    def _conflicts(self, doc: Document, analysis: AnalysisResult):
        """Build conflicts section."""
        self._section_heading(doc, "Detected Conflicts")

        severity_colors = {"HIGH": AMD_RED, "MEDIUM": AMBER_600, "LOW": GREEN_600}

        for i, conflict in enumerate(analysis.conflicts, 1):
            # Title
            para = doc.add_paragraph()
            r = para.add_run(f"#{i} {conflict.type}")
            r.bold = True
            r.font.size = Pt(11)
            r2 = para.add_run(f"  [{conflict.severity}]")
            r2.bold = True
            r2.font.color.rgb = severity_colors.get(conflict.severity, SLATE_700)

            # Excerpts table
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"

            # Doc names
            for j, d in enumerate([conflict.documentA, conflict.documentB]):
                cell = table.rows[0].cells[j]
                cell.text = ""
                _set_cell_bg(cell, "F1F5F9")
                run = cell.paragraphs[0].add_run(d.name)
                run.bold = True
                run.font.size = Pt(9)

            # Excerpts
            for j, d in enumerate([conflict.documentA, conflict.documentB]):
                cell = table.rows[1].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(f'"{d.excerpt}"')
                run.font.size = Pt(9)
                run.italic = True

            doc.add_paragraph()

            # Explanation
            exp = doc.add_paragraph()
            exp.add_run("Explanation: ").bold = True
            exp.add_run(conflict.explanation)

            # Action
            act = doc.add_paragraph()
            act.add_run("Recommended Action: ").bold = True
            r = act.add_run(conflict.recommendedAction)
            r.font.color.rgb = BLUE_700

            doc.add_paragraph()

    def _recommendation(self, doc: Document, analysis: AnalysisResult):
        """Build recommendation section."""
        self._section_heading(doc, "AI Recommendation")
        rec = analysis.recommendation

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(rec.title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = SLATE_900

        doc.add_paragraph()

        # Confidence
        conf_pct = int(rec.confidence * 100)
        conf_color = GREEN_600 if conf_pct >= 70 else AMBER_600 if conf_pct >= 40 else AMD_RED
        conf_para = doc.add_paragraph()
        conf_para.add_run("AI Confidence:  ").bold = True
        r = conf_para.add_run(f"{conf_pct}%")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = conf_color

        doc.add_paragraph()

        # Summary
        doc.add_paragraph(rec.summary).paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        # Next Steps
        if rec.nextSteps:
            ns = doc.add_paragraph()
            ns.add_run("Next Steps").bold = True

            for idx, step in enumerate(rec.nextSteps, 1):
                step_para = doc.add_paragraph()
                step_para.paragraph_format.left_indent = Cm(0.5)
                r = step_para.add_run(f"{idx}. ")
                r.bold = True
                r.font.color.rgb = AMD_RED
                step_para.add_run(step)

        doc.add_paragraph()

    def _footer(self, doc: Document):
        """Build document footer."""
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Generated by Clausify AI  |  Powered by AMD Instinct MI300X  |  clausify.app")
        run.font.size = Pt(9)
        run.font.color.rgb = SLATE_500
